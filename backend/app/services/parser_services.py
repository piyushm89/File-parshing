import fitz
from docx import Document
from openpyxl import load_workbook


def extract_text_from_pdf(file_path: str) -> str:
    pdf = fitz.open(file_path)
    parts = [page.get_text() for page in pdf]
    pdf.close()
    return "".join(parts)


def extract_text_from_docx(file_path: str) -> str:
    document = Document(file_path)
    parts = []

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            parts.append(text + "\n")

    for table in document.tables:
        for row in table.rows:
            row_data = []
            for cell in row.cells:
                value = cell.text.strip()
                if value:
                    row_data.append(value)
            if row_data:
                parts.append(" : ".join(row_data) + "\n")

    return "".join(parts)


def extract_text_from_excel(file_path: str) -> str:
    workbook = load_workbook(file_path, data_only=True)
    parts = []

    for sheet in workbook.worksheets:
        for row in sheet.iter_rows(values_only=True):
            row_text = " | ".join(
                str(cell) for cell in row if cell is not None
            )
            if row_text:
                parts.append(row_text + "\n")

    return "".join(parts)
